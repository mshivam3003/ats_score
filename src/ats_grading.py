import io
import json
import os
import re
import tempfile
from pathlib import Path
from typing import Any

from dotenv import load_dotenv
from docx import Document
from openai import OpenAI
from PyPDF2 import PdfReader

try:
    import pypandoc
except ImportError:
    pypandoc = None

MAX_INPUT_CHARS = 28000
SUPPORTED_EXTENSIONS = {
    ".pdf": "PDF",
    ".txt": "Plain text",
    ".docx": "Microsoft Word (.docx)",
    ".doc": "Microsoft Word (.doc)",
}
OPENROUTER_MODEL = os.getenv("OPENROUTER_MODEL", "google/gemini-2.5-flash-lite")
OPENROUTER_BASE_URL = os.getenv("OPENROUTER_BASE_URL", "https://openrouter.ai/api/v1")
OPENROUTER_MAX_TOKENS = int(os.getenv("OPENROUTER_MAX_TOKENS", "0") or 0)
OPENROUTER_OPTIMIZE_MAX_TOKENS = int(os.getenv("OPENROUTER_OPTIMIZE_MAX_TOKENS", "900") or 900)
MAX_REPAIR_CHARS = 6000


def get_api_key() -> str | None:
    load_dotenv()
    for key in ("OPENROUTER_API_KEY", "OPENAI_API_KEY", "OPENAI_APIKEY", "HF_TOKEN"):
        value = os.getenv(key)
        if value:
            return value
    return None


def get_openrouter_client():
    api_key = get_api_key()
    if not api_key:
        raise RuntimeError("OpenRouter API key not found. Please set OPENROUTER_API_KEY in .env.")

    return OpenAI(
        base_url=OPENROUTER_BASE_URL,
        api_key=api_key,
    )


def safe_truncate(text: str, max_chars: int = MAX_INPUT_CHARS) -> str:
    if len(text) <= max_chars:
        return text
    return text[:max_chars].rstrip() + "\n\n[...TEXT TRUNCATED DUE TO SIZE...]"


def extract_text_from_docx(file_bytes: bytes) -> str:
    document = Document(io.BytesIO(file_bytes))
    paragraphs = [para.text for para in document.paragraphs if para.text.strip()]
    table_text = []
    for table in document.tables:
        for row in table.rows:
            row_text = " \t ".join(cell.text for cell in row.cells if cell.text.strip())
            if row_text:
                table_text.append(row_text)
    return "\n".join(paragraphs + table_text).strip()


def extract_text_from_pdf(file_bytes: bytes) -> str:
    reader = PdfReader(io.BytesIO(file_bytes))
    pages = []
    for page in reader.pages:
        page_text = page.extract_text() or ""
        if page_text.strip():
            pages.append(page_text)
    return "\n\n".join(pages).strip()


def extract_text_from_doc(file_bytes: bytes) -> str:
    if pypandoc is None:
        raise RuntimeError(
            "DOC support requires pypandoc. Install it or upload a .docx / .pdf / .txt file instead."
        )

    with tempfile.NamedTemporaryFile(delete=False, suffix=".doc") as tmp:
        tmp.write(file_bytes)
        tmp.flush()
        temp_path = tmp.name

    try:
        return pypandoc.convert_file(temp_path, "plain") or ""
    finally:
        try:
            os.remove(temp_path)
        except OSError:
            pass


def extract_resume_text(uploaded_file) -> str:
    if uploaded_file is None:
        return ""

    extension = Path(uploaded_file.name).suffix.lower()
    file_bytes = uploaded_file.read()

    if extension == ".txt":
        return file_bytes.decode("utf-8", errors="ignore").strip()
    if extension == ".pdf":
        return extract_text_from_pdf(file_bytes)
    if extension == ".docx":
        return extract_text_from_docx(file_bytes)
    if extension == ".doc":
        return extract_text_from_doc(file_bytes)

    raise ValueError(
        f"Unsupported resume format: {extension}. Supported formats: {', '.join(SUPPORTED_EXTENSIONS)}"
    )


def _extract_affordable_max_tokens(error_text: str) -> int | None:
    match = re.search(r"can only afford\s+(\d+)", error_text, re.IGNORECASE)
    if not match:
        return None
    try:
        return int(match.group(1))
    except ValueError:
        return None


def _looks_like_openrouter_credit_error(exc: Exception) -> bool:
    text = str(exc)
    return "Error code: 402" in text or "requires more credits" in text or "'code': 402" in text


def make_chat_completion(client, prompt: str, temperature: float = 0.25, max_tokens: int = 1200):
    requested_max_tokens = int(max_tokens)
    if OPENROUTER_MAX_TOKENS > 0:
        requested_max_tokens = min(requested_max_tokens, OPENROUTER_MAX_TOKENS)

    def _call(max_tokens_to_use: int):
        return client.chat.completions.create(
            model=OPENROUTER_MODEL,
            messages=[{"role": "user", "content": prompt}],
            temperature=temperature,
            max_tokens=max_tokens_to_use,
            extra_body={"reasoning": {"effort": "low"}},
        )

    try:
        return _call(requested_max_tokens)
    except Exception as exc:
        if not _looks_like_openrouter_credit_error(exc):
            raise

        affordable = _extract_affordable_max_tokens(str(exc))
        if affordable is not None:
            retry_max = min(requested_max_tokens, max(128, affordable - 64))
        else:
            retry_max = max(128, int(requested_max_tokens * 0.75))

        if retry_max >= requested_max_tokens:
            raise RuntimeError(
                f"OpenRouter credit limit: lower max_tokens (set OPENROUTER_MAX_TOKENS). Original error: {exc}"
            ) from exc

        try:
            return _call(retry_max)
        except Exception as exc2:
            if _looks_like_openrouter_credit_error(exc2):
                raise RuntimeError(
                    f"OpenRouter credit limit even after lowering max_tokens to {retry_max}. "
                    f"Set OPENROUTER_MAX_TOKENS lower or add credits. Original error: {exc2}"
                ) from exc2
            raise


def _strip_code_fences(text: str) -> str:
    text = text.strip()
    if not text.startswith("```"):
        return text

    text = re.sub(r"^\s*```(?:json)?\s*", "", text, flags=re.IGNORECASE)
    text = re.sub(r"\s*```\s*$", "", text)
    return text.strip()


def _raw_decode_json(text: str) -> Any:
    text = _strip_code_fences(text)
    if not text:
        raise ValueError("LLM output was empty.")

    decoder = json.JSONDecoder()

    direct = text.lstrip()
    if direct.startswith("{") or direct.startswith("["):
        value, _end = decoder.raw_decode(direct)
        return value

    obj_index = text.find("{")
    arr_index = text.find("[")
    candidates = [i for i in (obj_index, arr_index) if i != -1]
    if not candidates:
        raise ValueError("LLM output contained no JSON object/array.")

    start = min(candidates)
    value, _end = decoder.raw_decode(text[start:])
    return value


def parse_llm_json_dict(content: str) -> dict:
    try:
        value = _raw_decode_json(content)
    except json.JSONDecodeError as exc:
        raise ValueError("LLM output was not valid JSON.") from exc

    if not isinstance(value, dict):
        raise ValueError("LLM output JSON was not an object.")
    return value


def _ensure_list(value: Any) -> list:
    if value is None:
        return []
    if isinstance(value, list):
        return value
    if isinstance(value, (set, tuple)):
        return list(value)
    if isinstance(value, str):
        return [value] if value.strip() else []
    return []


def _ensure_bool(value: Any) -> bool:
    if isinstance(value, bool):
        return value
    if isinstance(value, (int, float)):
        return bool(value)
    if isinstance(value, str):
        normalized = value.strip().lower()
        if normalized in {"true", "yes", "y", "1"}:
            return True
        if normalized in {"false", "no", "n", "0"}:
            return False
    return False


def _ensure_number(value: Any, default: float = 0.0) -> float:
    if isinstance(value, (int, float)):
        return float(value)
    if isinstance(value, str):
        match = re.search(r"-?\d+(?:\.\d+)?", value)
        if match:
            try:
                return float(match.group())
            except ValueError:
                return default
    return default


def normalize_assess_payload(parsed: dict) -> dict:
    sections = parsed.get("sections", [])
    if isinstance(sections, dict):
        sections = list(sections.keys())
    sections_list = [str(item).strip().lower() for item in _ensure_list(sections) if str(item).strip()]

    return {
        "skills": [str(item).strip() for item in _ensure_list(parsed.get("skills")) if str(item).strip()],
        "action_verbs": [str(item).strip() for item in _ensure_list(parsed.get("action_verbs")) if str(item).strip()],
        "sections": sections_list,
        "experience_years": _ensure_number(parsed.get("experience_years"), default=0.0),
        "has_metrics": _ensure_bool(parsed.get("has_metrics")),
        "keywords_from_jd": [str(item).strip() for item in _ensure_list(parsed.get("keywords_from_jd")) if str(item).strip()],
        "matched_keywords": [str(item).strip() for item in _ensure_list(parsed.get("matched_keywords")) if str(item).strip()],
        "recommendations": [str(item).strip() for item in _ensure_list(parsed.get("recommendations")) if str(item).strip()],
        "ats_rubric": parsed.get("ats_rubric") if isinstance(parsed.get("ats_rubric"), dict) else {},
        "role_title": str(parsed.get("role_title") or "").strip(),
        "target_roles": [str(item).strip() for item in _ensure_list(parsed.get("target_roles")) if str(item).strip()],
    }


def normalize_roles_payload(parsed: dict) -> dict:
    return {
        "best_role": str(parsed.get("best_role") or "").strip(),
        "target_roles": [str(item).strip() for item in _ensure_list(parsed.get("target_roles")) if str(item).strip()],
        "role_summary": str(parsed.get("role_summary") or "").strip(),
    }


def _dedupe_preserve_order(items: list[str]) -> list[str]:
    seen: set[str] = set()
    out: list[str] = []
    for item in items:
        key = (item or "").strip().lower()
        if not key:
            continue
        if key in seen:
            continue
        seen.add(key)
        out.append(item.strip())
    return out


def _compute_matched_keywords(resume_text: str, keywords_from_jd: list[str]) -> list[str]:
    resume_lower = (resume_text or "").lower()
    matched: list[str] = []
    seen: set[str] = set()
    for keyword in keywords_from_jd:
        keyword_clean = (keyword or "").strip()
        if not keyword_clean:
            continue
        if len(keyword_clean) < 3:
            continue
        key_lower = keyword_clean.lower()
        if key_lower in seen:
            continue
        if key_lower in resume_lower:
            matched.append(keyword_clean)
            seen.add(key_lower)
    return matched


def fallback_ats_rubric() -> dict:
    return {
        "parsability": {
            "max_points": 45,
            "signals": ["experience section", "education section", "skills section", "simple formatting", "consistent headings"],
        },
        "relevance": {
            "max_points": 55,
            "signals": ["JD keyword coverage", "hard skills alignment", "relevant experience", "action verbs", "quantified impact"],
        },
        "notes": "Fallback rubric used because the LLM did not return a rubric.",
    }


def fallback_recommendations(parsed: dict) -> list[str]:
    recommendations: list[str] = []

    sections = set((parsed.get("sections") or []))
    if "experience" not in sections:
        recommendations.append("Add a clearly labeled Experience section with role, company, dates, and 3-6 impact bullets per job.")
    if "education" not in sections:
        recommendations.append("Add an Education section with degree, institution, year, and relevant coursework (if applicable).")
    if "skills" not in sections:
        recommendations.append("Add a dedicated Skills section and group skills by category (Languages, Frameworks, Tools, Cloud, Databases).")

    missing = list(set(parsed.get("keywords_from_jd", [])) - set(parsed.get("matched_keywords", [])))
    if missing:
        sample = ", ".join(missing[:12])
        recommendations.append(f"Add missing JD keywords where truthful (e.g., {sample}).")

    if not parsed.get("has_metrics"):
        recommendations.append("Add measurable impact to bullets (%, $, time saved, scale, latency, conversion, revenue).")

    recommendations.append("Rewrite bullets to start with strong action verbs and include scope + outcome (what you did, how, and result).")
    recommendations.append("Mirror the JD phrasing for core skills/tools (without keyword stuffing) to improve ATS relevance.")

    deduped: list[str] = []
    seen: set[str] = set()
    for item in recommendations:
        if item not in seen:
            deduped.append(item)
            seen.add(item)
    return deduped[:10]


def llm_assess_resume(resume_text: str, jd_text: str, industry: str, experience_level: str, client):
    expected_keys = [
        "skills",
        "action_verbs",
        "sections",
        "experience_years",
        "has_metrics",
        "keywords_from_jd",
        "matched_keywords",
        "recommendations",
        "ats_rubric",
        "role_title",
        "target_roles",
    ]

    prompt = f"""
    You are an expert ATS resume analyst.
    Read the resume and the job description and extract structured information needed for ATS scoring.

    Resume:
    {resume_text}

    Job Description:
    {jd_text}

    Industry: {industry}
    Experience Level: {experience_level}

    Output rules:
    - Return ONLY strict JSON (no markdown, no code fences, no commentary).
    - The top-level must be a JSON object.
    - Use these types:
      - skills: array of strings
      - action_verbs: array of strings
      - sections: array of strings (lowercase section names like "experience", "education", "skills")
      - experience_years: number
      - has_metrics: boolean
      - keywords_from_jd: array of strings
      - matched_keywords: array of strings (subset of keywords_from_jd)
      - recommendations: array of strings
      - ats_rubric: object
      - role_title: string
      - target_roles: array of strings
    - Must include all keys: {", ".join(expected_keys)}
    """

    response = make_chat_completion(client, prompt)
    content = getattr(response.choices[0].message, "content", None)
    try:
        parsed = parse_llm_json_dict(content or "")
    except ValueError:
        repair_prompt = f"""
        Your previous output was not valid JSON. Convert it into ONE strict JSON object.
        Rules:
        - Output ONLY JSON. No markdown. No code fences.
        - Must include keys: {", ".join(expected_keys)}
        - If a value is unknown, use [] for arrays, {{}} for objects, "" for strings, 0 for numbers, false for booleans.

        Previous output:
        {safe_truncate(content or "", MAX_REPAIR_CHARS)}
        """
        repaired = make_chat_completion(client, repair_prompt, temperature=0.0, max_tokens=1400)
        repaired_content = getattr(repaired.choices[0].message, "content", None)
        parsed = parse_llm_json_dict(repaired_content or "")

    normalized = normalize_assess_payload(parsed)
    normalized["keywords_from_jd"] = _dedupe_preserve_order(normalized.get("keywords_from_jd", []))[:60]
    if normalized["keywords_from_jd"]:
        # Deterministic match to reduce scoring volatility.
        normalized["matched_keywords"] = _compute_matched_keywords(resume_text, normalized["keywords_from_jd"])

    needs_fill = (
        not normalized["keywords_from_jd"]
        or not normalized["recommendations"]
        or not normalized["ats_rubric"]
    )
    if needs_fill:
        fill_prompt = f"""
        You previously returned JSON with missing/empty fields. Return ONE strict JSON object that includes ALL keys.
        Requirements:
        - keywords_from_jd: extract 15-40 important JD keywords/phrases (strings).
        - matched_keywords: subset of keywords_from_jd that appear in the resume text.
        - recommendations: 8-12 actionable, specific resume improvements tailored to the JD.
        - ats_rubric: a non-empty object describing how you scored parsability and relevance.
        - Keep other fields reasonable and consistent.

        Resume:
        {resume_text}

        Job Description:
        {jd_text}

        Industry: {industry}
        Experience Level: {experience_level}

        Current JSON (may contain empties):
        {json.dumps(normalized, ensure_ascii=False)}
        """
        filled = make_chat_completion(client, fill_prompt, temperature=0.25, max_tokens=1600)
        filled_content = getattr(filled.choices[0].message, "content", None)
        try:
            filled_parsed = parse_llm_json_dict(filled_content or "")
            filled_normalized = normalize_assess_payload(filled_parsed)
            merged = {**normalized, **filled_normalized}
            merged["keywords_from_jd"] = _dedupe_preserve_order(merged.get("keywords_from_jd", []))[:60]
            if merged["keywords_from_jd"]:
                merged["matched_keywords"] = _compute_matched_keywords(resume_text, merged["keywords_from_jd"])
            normalized = merged
        except ValueError:
            pass

    return normalized


def llm_optimize_resume(
    resume_text: str,
    jd_text: str,
    recommendations: list[str],
    ats_rubric: dict,
    missing_keywords: list[str],
    client,
):
    prompt = f"""
    You are an AI resume optimizer. Improve ATS compatibility and JD alignment.

    Hard rules:
    - DO NOT delete or omit any existing content from the resume.
    - Preserve all roles, companies, projects, dates, and achievements; you may rewrite for clarity and ATS alignment.
    - Keep the resume structure and section order. If you add content, add it as new lines; do not remove lines.
    - If any of the missing keywords apply truthfully, incorporate them naturally into Skills/Summary/Experience bullets.

    Original Resume:
    {resume_text}

    Job Description:
    {jd_text}

    Recommendations:
    {json.dumps(recommendations, indent=2)}

    Missing JD keywords (include only if truthful):
    {json.dumps(_dedupe_preserve_order([str(k).strip() for k in (missing_keywords or []) if str(k).strip()])[:25], indent=2)}

    ATS scoring rubric:
    {json.dumps(ats_rubric, indent=2)}

    Return ONLY the complete optimized resume text in plain text format (no markdown, no code fences).
    """

    response = make_chat_completion(
        client,
        prompt,
        temperature=0.35,
        max_tokens=OPENROUTER_OPTIMIZE_MAX_TOKENS,
    )
    optimized = (response.choices[0].message.content or "").strip()

    # If the model output is suspiciously short, fall back to an "augment-only" optimizer
    # that never drops content: add an ATS-optimized Summary/Skills header, then include the full resume.
    min_expected = max(900, int(len(resume_text) * 0.65)) if resume_text else 0
    if optimized and len(optimized) >= min_expected:
        return optimized

    augment_prompt = f"""
    You are an AI resume optimizer.
    Create ONLY these two sections to prepend to the resume (plain text):
    1) PROFESSIONAL SUMMARY (4-6 lines tailored to the JD)
    2) CORE SKILLS (grouped, ATS-friendly keywords; 3-6 groups)

    Rules:
    - Output only those sections in plain text (no markdown, no code fences).
    - Do not include the rest of the resume.

    Resume:
    {resume_text}

    Job Description:
    {jd_text}
    """
    header = make_chat_completion(client, augment_prompt, temperature=0.35, max_tokens=450)
    header_text = (header.choices[0].message.content or "").strip()
    header_text = header_text.strip()

    if header_text:
        return f"{header_text}\n\n{resume_text.strip()}"
    return resume_text.strip()


def llm_infer_target_roles(resume_text: str, client):
    expected_keys = ["best_role", "target_roles", "role_summary"]
    prompt = f"""
    You are a career intelligence engine. Read the resume and identify the one best-fit job title and a short list of target job titles.

    Resume:
    {resume_text}

    Output rules:
    - Return ONLY strict JSON (no markdown, no code fences, no commentary).
    - The top-level must be a JSON object.
    - Must include keys: {", ".join(expected_keys)}
    """

    response = make_chat_completion(client, prompt)
    content = getattr(response.choices[0].message, "content", None)
    try:
        parsed = parse_llm_json_dict(content or "")
    except ValueError:
        repair_prompt = f"""
        Your previous output was not valid JSON. Convert it into ONE strict JSON object.
        Rules:
        - Output ONLY JSON. No markdown. No code fences.
        - Must include keys: {", ".join(expected_keys)}
        - If a value is unknown, use [] for arrays, "" for strings.

        Previous output:
        {safe_truncate(content or "", MAX_REPAIR_CHARS)}
        """
        repaired = make_chat_completion(client, repair_prompt, temperature=0.0, max_tokens=600)
        repaired_content = getattr(repaired.choices[0].message, "content", None)
        parsed = parse_llm_json_dict(repaired_content or "")

    return normalize_roles_payload(parsed)


def parsability_score(parsed):
    score = 45
    if "experience" not in parsed["sections"]:
        score -= 10
    if "education" not in parsed["sections"]:
        score -= 5
    if "skills" not in parsed["sections"]:
        score -= 5
    return max(score, 0)


def keyword_score(parsed):
    total = len(parsed["keywords_from_jd"])
    matched = len(parsed["matched_keywords"])
    if total == 0:
        return 0
    return (matched / total) * 25


def skill_score(parsed):
    return min(len(parsed["skills"]) * 1.5, 10)


def experience_score(parsed):
    exp = parsed["experience_years"]
    if exp >= 3:
        return 10
    elif exp >= 1:
        return 7
    return 4


def action_score(parsed):
    return min(len(parsed["action_verbs"]) * 1.5, 5)


def metric_score(parsed):
    return 5 if parsed["has_metrics"] else 1


def relevance_score(parsed):
    return (
        keyword_score(parsed)
        + skill_score(parsed)
        + experience_score(parsed)
        + action_score(parsed)
        + metric_score(parsed)
    )


def final_ats_score(parsed):
    p_score = parsability_score(parsed)
    r_score = relevance_score(parsed)
    final = p_score + r_score
    final = max(0, min(100, int(final)))
    if final >= 80:
        match_level = "Strong Match"
    elif final >= 60:
        match_level = "Moderate Match"
    else:
        match_level = "Low Match"
    jd_total = len(parsed.get("keywords_from_jd", []))
    matched_total = len(parsed.get("matched_keywords", []))
    missing_keywords = list(set(parsed.get("keywords_from_jd", [])) - set(parsed.get("matched_keywords", []))) if jd_total else []
    return {
        "ATS_score": final,
        "parsability_score": int(p_score),
        "relevance_score": int(r_score),
        "match_level": match_level,
        "missing_keywords": missing_keywords,
        "jd_keywords_total": jd_total,
        "matched_keywords_total": matched_total,
        "skills_detected": parsed.get("skills", []),
    }


def run_ats_pipeline(
    resume_text: str,
    jd_text: str,
    industry: str,
    experience_level: str,
    client,
    jd_keywords_override: list[str] | None = None,
):
    parsed = llm_assess_resume(resume_text, jd_text, industry, experience_level, client)
    # Keyword extraction can be noisy across calls; optionally freeze JD keyword set to compare before/after fairly.
    if jd_keywords_override:
        parsed["keywords_from_jd"] = _dedupe_preserve_order([str(k).strip() for k in jd_keywords_override if str(k).strip()])[:60]
    else:
        parsed["keywords_from_jd"] = _dedupe_preserve_order(parsed.get("keywords_from_jd", []))[:60]

    if parsed.get("keywords_from_jd"):
        parsed["matched_keywords"] = _compute_matched_keywords(resume_text, parsed["keywords_from_jd"])
    result = final_ats_score(parsed)
    result["recommendations"] = parsed.get("recommendations", []) or fallback_recommendations(parsed)
    result["ats_rubric"] = parsed.get("ats_rubric", {}) or fallback_ats_rubric()
    result["role_title"] = parsed.get("role_title", "")
    result["target_roles"] = parsed.get("target_roles", [])
    result["jd_keywords"] = parsed.get("keywords_from_jd", [])
    result["matched_keywords"] = parsed.get("matched_keywords", [])
    return result


def create_resume_docx_bytes(text: str) -> bytes:
    document = Document()
    for line in text.splitlines():
        document.add_paragraph(line)
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def create_resume_pdf_bytes(docx_bytes: bytes) -> bytes:
    if pypandoc is None:
        raise RuntimeError("PDF export requires pypandoc. Install it to enable PDF generation.")

    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_docx:
        tmp_docx.write(docx_bytes)
        tmp_docx.flush()
        docx_path = tmp_docx.name

    pdf_path = docx_path.replace(".docx", ".pdf")
    try:
        pypandoc.convert_file(docx_path, "pdf", outputfile=pdf_path)
        with open(pdf_path, "rb") as f:
            return f.read()
    finally:
        for path in (docx_path, pdf_path):
            try:
                os.remove(path)
            except OSError:
                pass


def _pdf_literal_string(text: str) -> str:
    raw = (text or "").encode("cp1252", errors="replace")
    out: list[str] = []
    for byte in raw:
        if byte in (0x28, 0x29, 0x5C):  # ( ) \
            out.append("\\" + chr(byte))
        elif 32 <= byte <= 126:
            out.append(chr(byte))
        else:
            out.append(f"\\{byte:03o}")
    return "(" + "".join(out) + ")"


def create_resume_pdf_from_text_bytes(text: str) -> bytes:
    """
    Pure-Python, dependency-free PDF export from plain text.
    This is a fallback when `pypandoc`/Pandoc are not available.
    """
    import textwrap

    page_width = 612
    page_height = 792
    margin_left = 72
    margin_top = 72
    font_size = 11
    leading = 14
    max_chars_per_line = 95

    lines: list[str] = []
    for raw_line in (text or "").splitlines():
        if not raw_line.strip():
            lines.append("")
            continue
        wrapped = textwrap.wrap(raw_line, width=max_chars_per_line, break_long_words=False, break_on_hyphens=False)
        lines.extend(wrapped if wrapped else [""])

    lines_per_page = max(1, int((page_height - margin_top * 2) / leading))
    pages: list[list[str]] = [lines[i : i + lines_per_page] for i in range(0, len(lines), lines_per_page)] or [[]]

    objects: list[bytes] = []

    # 1) Catalog, 2) Pages, 3) Font
    objects.append(b"<< /Type /Catalog /Pages 2 0 R >>")
    objects.append(b"")  # placeholder for Pages obj
    objects.append(b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>")

    page_obj_numbers: list[int] = []
    content_obj_numbers: list[int] = []

    next_obj_num = 4
    for _page in pages:
        page_obj_numbers.append(next_obj_num)
        next_obj_num += 1
        content_obj_numbers.append(next_obj_num)
        next_obj_num += 1

    # Build content streams and page objects
    for page_index, page_lines in enumerate(pages):
        content_lines: list[str] = []
        content_lines.append("BT")
        content_lines.append(f"/F1 {font_size} Tf")
        content_lines.append(f"{leading} TL")
        start_x = margin_left
        start_y = page_height - margin_top
        content_lines.append(f"{start_x} {start_y} Td")

        first = True
        for line in page_lines:
            if not first:
                content_lines.append("T*")
            first = False
            content_lines.append(f"{_pdf_literal_string(line)} Tj")
        content_lines.append("ET")

        stream_data = ("\n".join(content_lines) + "\n").encode("ascii", errors="ignore")
        content_obj = (
            b"<< /Length "
            + str(len(stream_data)).encode("ascii")
            + b" >>\nstream\n"
            + stream_data
            + b"endstream"
        )

        content_obj_num = content_obj_numbers[page_index]
        page_obj_num = page_obj_numbers[page_index]
        objects.append(b"")  # placeholder until we add objects at exact positions
        objects.append(b"")

        # Place them by object number (objects list is 1-indexed conceptually)
        # We'll just set them by index later in a stable way.
        objects_index_for_page = page_obj_num - 1
        objects_index_for_content = content_obj_num - 1
        while len(objects) <= max(objects_index_for_page, objects_index_for_content):
            objects.append(b"")

        page_obj = (
            b"<< /Type /Page /Parent 2 0 R "
            b"/MediaBox [0 0 612 792] "
            b"/Resources << /Font << /F1 3 0 R >> >> "
            b"/Contents "
            + str(content_obj_num).encode("ascii")
            + b" 0 R >>"
        )
        objects[objects_index_for_page] = page_obj
        objects[objects_index_for_content] = content_obj

    kids = " ".join(f"{n} 0 R" for n in page_obj_numbers).encode("ascii")
    pages_obj = b"<< /Type /Pages /Kids [ " + kids + b" ] /Count " + str(len(page_obj_numbers)).encode("ascii") + b" >>"
    objects[1] = pages_obj

    # Assemble file with xref
    header = b"%PDF-1.4\n%\xe2\xe3\xcf\xd3\n"
    body_parts: list[bytes] = [header]
    offsets = [0] * (len(objects) + 1)

    current_offset = len(header)
    for obj_num, obj_content in enumerate(objects, start=1):
        offsets[obj_num] = current_offset
        obj_bytes = (
            str(obj_num).encode("ascii")
            + b" 0 obj\n"
            + obj_content
            + b"\nendobj\n"
        )
        body_parts.append(obj_bytes)
        current_offset += len(obj_bytes)

    xref_offset = current_offset
    xref_lines = [b"xref\n", f"0 {len(objects)+1}\n".encode("ascii")]
    xref_lines.append(b"0000000000 65535 f \n")
    for obj_num in range(1, len(objects) + 1):
        xref_lines.append(f"{offsets[obj_num]:010d} 00000 n \n".encode("ascii"))

    trailer = (
        b"trailer\n<< /Size "
        + str(len(objects) + 1).encode("ascii")
        + b" /Root 1 0 R >>\nstartxref\n"
        + str(xref_offset).encode("ascii")
        + b"\n%%EOF\n"
    )

    return b"".join(body_parts + xref_lines + [trailer])
